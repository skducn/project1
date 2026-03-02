from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn

def create_medical_disease_lib_docs():
    # 1. 创建文档对象
    doc = Document()
    # 设置页面边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.8)

    # 2. 定义全局样式（宋体、小四、1.5倍行间距，符合模版要求）
    # 定义正文样式
    normal_style = doc.styles['Normal']
    normal_style.font.name = '宋体'
    normal_style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    normal_style.font.size = Pt(12)  # 小四=12pt
    normal_style.paragraph_format.line_spacing = 1.5
    normal_style.paragraph_format.first_line_indent = Cm(0.74)  # 首行缩进2字符
    normal_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # 定义标题1样式（一级标题，居中/左对齐，加粗）
    h1_style = doc.styles.add_style('Heading 1 Custom', WD_STYLE_TYPE.PARAGRAPH)
    h1_style.font.name = '宋体'
    h1_style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    h1_style.font.size = Pt(16)
    h1_style.font.bold = True
    h1_style.paragraph_format.line_spacing = 1.5
    h1_style.paragraph_format.space_after = Pt(18)
    h1_style.paragraph_format.space_before = Pt(18)

    # 定义标题2样式（二级标题）
    h2_style = doc.styles.add_style('Heading 2 Custom', WD_STYLE_TYPE.PARAGRAPH)
    h2_style.font.name = '宋体'
    h2_style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    h2_style.font.size = Pt(14)
    h2_style.font.bold = True
    h2_style.paragraph_format.line_spacing = 1.5
    h2_style.paragraph_format.space_after = Pt(12)
    h2_style.paragraph_format.space_before = Pt(12)

    # 定义标题3样式（三级标题）
    h3_style = doc.styles.add_style('Heading 3 Custom', WD_STYLE_TYPE.PARAGRAPH)
    h3_style.font.name = '宋体'
    h3_style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    h3_style.font.size = Pt(13)
    h3_style.font.bold = True
    h3_style.paragraph_format.line_spacing = 1.5

    # 定义标题4样式（四级标题）
    h4_style = doc.styles.add_style('Heading 4 Custom', WD_STYLE_TYPE.PARAGRAPH)
    h4_style.font.name = '宋体'
    h4_style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    h4_style.font.size = Pt(12)
    h4_style.font.bold = True
    h4_style.paragraph_format.line_spacing = 1.5

    # 定义无首行缩进样式（用于列表、表格说明、封面等）
    no_indent = doc.styles.add_style('No Indent', WD_STYLE_TYPE.PARAGRAPH)
    no_indent.font.name = '宋体'
    no_indent._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    no_indent.font.size = Pt(12)
    no_indent.paragraph_format.line_spacing = 1.5
    no_indent.paragraph_format.first_line_indent = Cm(0)
    no_indent.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # 定义居中加粗样式（封面标题）
    center_bold = doc.styles.add_style('Center Bold', WD_STYLE_TYPE.PARAGRAPH)
    center_bold.font.name = '宋体'
    center_bold._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    center_bold.font.size = Pt(20)
    center_bold.font.bold = True
    center_bold.paragraph_format.line_spacing = 1.5
    center_bold.paragraph_format.first_line_indent = Cm(0)
    center_bold.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    center_bold.paragraph_format.space_after = Pt(24)

    # 定义居中样式（封面副标题）
    center_normal = doc.styles.add_style('Center Normal', WD_STYLE_TYPE.PARAGRAPH)
    center_normal.font.name = '宋体'
    center_normal._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    center_normal.font.size = Pt(14)
    center_normal.paragraph_format.line_spacing = 1.5
    center_normal.paragraph_format.first_line_indent = Cm(0)
    center_normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    center_normal.paragraph_format.space_after = Pt(36)

    # ======================================
    # 一、文档封面（必选，严格按模版格式）
    # ======================================
    title1 = doc.add_paragraph(style='Center Bold')
    title1.add_run('【专病库】需求规格说明书（测试用例生成专用）')

    title2 = doc.add_paragraph(style='Center Normal')
    title2.add_run('适配通义千问API，结构化需求提取')

    info = doc.add_paragraph(style='No Indent')
    info.add_run('版本号：V1.0 | 编制人：金浩 | 编制日期：2026-2-20')
    # 封面后换页
    doc.add_page_break()

    # ======================================
    # 二、文档目录（必选，模版要求自动生成，此处标注操作方式）
    # ======================================
    h1_2 = doc.add_paragraph('二、文档目录（必选，自动生成）', style='Heading 1 Custom')
    h1_2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    op = doc.add_paragraph(style='No Indent')
    op.add_run('操作方式：Word中“引用”→“目录”→“自动目录1”，确保所有一级/二级/三级标题自动收录，便于大模型定位各模块需求，避免遗漏。')
    doc.add_page_break()

    # ======================================
    # 三、核心需求总览（必选）
    # ======================================
    h1_3 = doc.add_paragraph('三、核心需求总览（必选，全局说明）', style='Heading 1 Custom')
    # 3.1 项目核心目标
    h2_31 = doc.add_paragraph('3.1 项目核心目标', style='Heading 2 Custom')
    p31 = doc.add_paragraph(style='Normal')
    p31.add_run('本项目为医疗专病库管理系统，核心实现专病信息录入、查询、编辑、删除、分类管理功能，保障专病数据的准确性、完整性和操作便捷性，本次测试覆盖所有核心功能模块。')
    # 3.2 测试范围说明
    h2_32 = doc.add_paragraph('3.2 测试范围说明（明确大模型生成用例的边界）', style='Heading 2 Custom')
    p32_1 = doc.add_paragraph(style='No Indent')
    p32_1.add_run('• 包含：所有核心功能模块（专病信息模块、分类管理模块、数据查询模块）的功能测试用例；')
    p32_2 = doc.add_paragraph(style='No Indent')
    p32_2.add_run('• 不包含：性能测试、安全测试、兼容性测试、接口测试相关用例；')
    p32_3 = doc.add_paragraph(style='No Indent')
    p32_3.add_run('• 特殊说明：仅覆盖正常操作、异常输入、边界值场景，不包含极端异常场景（如服务器宕机、数据库崩溃）。')
    # 3.3 通用约束
    h2_33 = doc.add_paragraph('3.3 通用约束（全局生效，所有模块共用）', style='Heading 2 Custom')
    p33_1 = doc.add_paragraph(style='No Indent')
    p33_1.add_run('• 输入约束：所有必填项为空时，无法提交操作，提示“请填写必填项”；')
    p33_2 = doc.add_paragraph(style='No Indent')
    p33_2.add_run('• 格式约束：日期格式统一为“XXXX-XX-XX”，编码为6位字母+数字组合，数值保留2位小数；')
    p33_3 = doc.add_paragraph(style='No Indent')
    p33_3.add_run('• 权限约束：未登录用户仅可访问专病库公开查询页面，无法进行录入、编辑、删除操作；')
    p33_4 = doc.add_paragraph(style='No Indent')
    p33_4.add_run('• 数据约束：所有提交的专病数据需经过格式校验，校验不通过则无法保存至数据库。')

    # ======================================
    # 四、模块详细需求（必选，核心结构化部分）
    # ======================================
    h1_4 = doc.add_paragraph('四、模块详细需求（必选，核心结构化部分）', style='Heading 1 Custom')
    p4 = doc.add_paragraph(style='No Indent')
    p4.add_run('核心要求：每个模块独立分区，一级标题为模块名称，二级标题为功能点，三级标题为功能细节（规则/约束/输入输出），字段固定，表述规范，通义千问可直接提取对应信息生成用例。')
    p4_1 = doc.add_paragraph(style='No Indent')
    p4_1.add_run('统一格式（每个模块均按此结构填写，不可修改标题层级和字段名称）：')

    # 4.1 专病信息模块
    h2_41 = doc.add_paragraph('4.1 专病信息模块', style='Heading 2 Custom')
    # 4.1.1 专病信息录入功能
    h3_411 = doc.add_paragraph('4.1.1 专病信息录入功能', style='Heading 3 Custom')
    # 4.1.1.1 功能描述
    h4_4111 = doc.add_paragraph('4.1.1.1 功能描述（必选）', style='Heading 4 Custom')
    p4111 = doc.add_paragraph(style='Normal')
    p4111.add_run('用户登录系统后，通过左侧导航栏“专病管理”→“新增专病”进入录入页面，填写专病编码、名称、所属科室、发病机制、诊疗方案等信息，点击“提交保存”完成录入，录入成功后自动跳转至专病信息列表页面。')
    # 4.1.1.2 输入项（带表格，模版固定格式）
    h4_4112 = doc.add_paragraph('4.1.1.2 输入项（必选，结构化列出，明确是否必填、格式、范围）', style='Heading 4 Custom')
    p4112 = doc.add_paragraph(style='No Indent')
    p4112.add_run('采用“输入项名称 | 是否必填 | 格式要求 | 取值范围 | 备注”的固定格式，表格如下：')
    # 创建输入项表格
    table1 = doc.add_table(rows=1, cols=5)
    table1.style = 'Table Grid'
    # 表格表头
    hdr_cells = table1.rows[0].cells
    hdr_cells[0].text = '输入项名称'
    hdr_cells[1].text = '是否必填'
    hdr_cells[2].text = '格式要求'
    hdr_cells[3].text = '取值范围'
    hdr_cells[4].text = '备注'
    # 表格内容（按业务填充）
    rows_data = [
        ['专病编码', '是', '字母+数字组合，无特殊字符', '6个字符', '系统唯一，不可重复'],
        ['专病名称', '是', '中文+英文，可含顿号/括号', '2-50个字符', '不可重复，需与临床标准一致'],
        ['所属科室', '是', '纯中文', '2-20个字符', '下拉框选择，不可手动输入'],
        ['发病机制', '否', '纯中文，可含标点', '0-500个字符', '简述核心发病原因'],
        ['诊疗方案', '是', '中文+数字+标点', '50-2000个字符', '需包含基础治疗、药物治疗'],
        ['发病率', '否', '数字+百分号', '0.00%-100.00%', '保留2位小数'],
        ['录入日期', '是', '日期格式', 'XXXX-XX-XX', '默认当前日期，可手动修改']
    ]
    for row_data in rows_data:
        row_cells = table1.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = val
    # 表格后换行
    doc.add_paragraph(style='No Indent')

    # 4.1.1.3 输出项
    h4_4113 = doc.add_paragraph('4.1.1.3 输出项（必选，明确操作后的反馈/结果）', style='Heading 4 Custom')
    p4113_1 = doc.add_paragraph(style='No Indent')
    p4113_1.add_run('• 正常场景：所有输入项符合要求，点击保存后提示“专病信息录入成功”，自动跳转至专病信息列表，数据库新增该专病完整数据；')
    p4113_2 = doc.add_paragraph(style='No Indent')
    p4113_2.add_run('• 异常场景1：专病编码重复，提交后提示“该专病编码已存在，请更换”，不跳转，不新增数据；')
    p4113_3 = doc.add_paragraph(style='No Indent')
    p4113_3.add_run('• 异常场景2：专病名称重复，提交后提示“该专病名称已录入，请勿重复添加”；')
    p4113_4 = doc.add_paragraph(style='No Indent')
    p4113_4.add_run('• 异常场景3：诊疗方案字符数不足50，提示“诊疗方案需至少填写50个字符，请补充”；')
    p4113_5 = doc.add_paragraph(style='No Indent')
    p4113_5.add_run('• 异常场景4：日期格式错误，提示“请输入正确的日期格式（XXXX-XX-XX）”。')

    # 4.1.1.4 边界场景
    h4_4114 = doc.add_paragraph('4.1.1.4 边界场景（可选，补充特殊情况，提升用例完整性）', style='Heading 4 Custom')
    p4114_1 = doc.add_paragraph(style='No Indent')
    p4114_1.add_run('• 专病编码取最小值6个字符（如：YB0001），可正常录入；')
    p4114_2 = doc.add_paragraph(style='No Indent')
    p4114_2.add_run('• 专病名称取最大值50个字符，可正常录入；')
    p4114_3 = doc.add_paragraph(style='No Indent')
    p4114_3.add_run('• 诊疗方案取最小值50个字符、最大值2000个字符，均可正常保存；')
    p4114_4 = doc.add_paragraph(style='No Indent')
    p4114_4.add_run('• 发病率取边界值0.00%、100.00%，格式校验通过可保存。')

    # 4.1.1.5 特殊规则
    h4_4115 = doc.add_paragraph('4.1.1.5 特殊规则（可选，补充该功能独有的规则，非全局约束）', style='Heading 4 Custom')
    p4115_1 = doc.add_paragraph(style='No Indent')
    p4115_1.add_run('• 同一用户10分钟内最多可录入5条专病信息，超过5条提示“操作过于频繁，请10分钟后再试”；')
    p4115_2 = doc.add_paragraph(style='No Indent')
    p4115_2.add_run('• 专病信息录入成功后，系统自动发送录入提醒至科室管理员邮箱；')
    p4115_3 = doc.add_paragraph(style='No Indent')
    p4115_3.add_run('• 录入的专病信息需经过管理员审核后，才可在公开查询页面展示。')

    # 4.1.2 专病信息编辑功能（沿用相同结构）
    h3_412 = doc.add_paragraph('4.1.2 专病信息编辑功能', style='Heading 3 Custom')
    h4_4121 = doc.add_paragraph('4.1.2.1 功能描述（必选）', style='Heading 4 Custom')
    p4121 = doc.add_paragraph(style='Normal')
    p4121.add_run('用户登录系统后，在专病信息列表页面选择需编辑的专病，点击“编辑”按钮进入编辑页面，可修改除专病编码外的所有信息，点击“保存修改”完成操作，修改成功后返回列表页面并刷新数据。')
    h4_4122 = doc.add_paragraph('4.1.2.2 输入项（必选，结构化列出，明确是否必填、格式、范围）', style='Heading 4 Custom')
    p4122 = doc.add_paragraph(style='No Indent')
    p4122.add_run('输入项格式与“4.1.1.2 输入项”一致，仅专病编码为不可编辑状态，其余项规则相同。')
    h4_4123 = doc.add_paragraph('4.1.2.3 输出项（必选，明确操作后的反馈/结果）', style='Heading 4 Custom')
    p4123_1 = doc.add_paragraph(style='No Indent')
    p4123_1.add_run('• 正常场景：修改内容符合规则，点击保存后提示“专病信息修改成功”，列表页面刷新为最新数据，数据库同步更新；')
    p4123_2 = doc.add_paragraph(style='No Indent')
    p4123_2.add_run('• 异常场景1：修改后专病名称与其他数据重复，提示“该专病名称已存在，修改失败”；')
    p4123_3 = doc.add_paragraph(style='No Indent')
    p4123_3.add_run('• 异常场景2：未修改任何内容直接点击保存，提示“未检测到修改内容，无需保存”。')
    h4_4124 = doc.add_paragraph('4.1.2.4 边界场景（可选）', style='Heading 4 Custom')
    p4124_1 = doc.add_paragraph(style='No Indent')
    p4124_1.add_run('• 将原最短专病名称修改为最大值50个字符，可正常保存；')
    p4124_2 = doc.add_paragraph(style='No Indent')
    p4124_2.add_run('• 将诊疗方案从50个字符修改为2000个字符，格式校验通过可保存。')
    h4_4125 = doc.add_paragraph('4.1.2.5 特殊规则（可选）', style='Heading 4 Custom')
    p4125_1 = doc.add_paragraph(style='No Indent')
    p4125_1.add_run('• 专病信息修改后，系统自动记录修改人、修改时间及修改前的原始数据，留存日志；')
    p4125_2 = doc.add_paragraph(style='No Indent')
    p4125_2.add_run('• 已审核通过的专病信息修改后，需重新提交管理员审核。')

    # 4.2 数据查询模块（简化填充，保持结构）
    h2_42 = doc.add_paragraph('4.2 数据查询模块', style='Heading 2 Custom')
    h3_421 = doc.add_paragraph('4.2.1 专病精准查询功能', style='Heading 3 Custom')
    h4_4211 = doc.add_paragraph('4.2.1.1 功能描述（必选）', style='Heading 4 Custom')
    p4211 = doc.add_paragraph(style='Normal')
    p4211.add_run('用户（登录/未登录）均可进入专病查询页面，输入专病编码或专病名称，点击“精准查询”按钮，系统匹配唯一数据并展示专病完整信息，未匹配到数据则提示无结果。')
    h4_4212 = doc.add_paragraph('4.2.1.2 输入项（必选）', style='Heading 4 Custom')
    p4212 = doc.add_paragraph(style='No Indent')
    p4212.add_run('采用“输入项名称 | 是否必填 | 格式要求 | 取值范围 | 备注”的固定格式，表格如下：')
    # 创建查询输入项表格
    table2 = doc.add_table(rows=1, cols=5)
    table2.style = 'Table Grid'
    hdr_cells2 = table2.rows[0].cells
    hdr_cells2[0].text = '输入项名称'
    hdr_cells2[1].text = '是否必填'
    hdr_cells2[2].text = '格式要求'
    hdr_cells2[3].text = '取值范围'
    hdr_cells2[4].text = '备注'
    row_data2 = [
        ['专病编码', '否', '字母+数字组合', '6个字符', '与编码录入格式一致'],
        ['专病名称', '否', '中文+英文，可含顿号/括号', '2-50个字符', '支持模糊匹配'],
        ['查询类型', '是', '纯中文', '2-10个字符', '下拉框选择：精准编码/精准名称']
    ]
    for row_data in row_data2:
        row_cells = table2.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = val
    doc.add_paragraph(style='No Indent')

    h4_4213 = doc.add_paragraph('4.2.1.3 输出项（必选）', style='Heading 4 Custom')
    p4213_1 = doc.add_paragraph(style='No Indent')
    p4213_1.add_run('• 正常场景：输入有效编码/名称，查询后展示专病编码、名称、所属科室、发病机制等完整信息；')
    p4213_2 = doc.add_paragraph(style='No Indent')
    p4213_2.add_run('• 异常场景1：输入不存在的编码/名称，提示“未查询到相关专病信息，请核对后重试”；')
    p4213_3 = doc.add_paragraph(style='No Indent')
    p4213_3.add_run('• 异常场景2：仅选择查询类型，未输入编码/名称，提示“请输入查询关键词后再提交”。')

    # ======================================
    # 五、异常场景汇总（可选）
    # ======================================
    h1_5 = doc.add_paragraph('五、异常场景汇总（可选，补充全局异常）', style='Heading 1 Custom')
    p5_1 = doc.add_paragraph(style='No Indent')
    p5_1.add_run('• 网络异常：所有功能操作时，若网络中断，提示“网络异常，请检查网络连接后重试”，操作不生效，未保存的数据不丢失；')
    p5_2 = doc.add_paragraph(style='No Indent')
    p5_2.add_run('• 系统异常：操作过程中系统报错，提示“系统繁忙，请稍后再试”，并自动记录错误日志至系统后台；')
    p5_3 = doc.add_paragraph(style='No Indent')
    p5_3.add_run('• 权限异常：未登录用户尝试点击录入/编辑/删除按钮，自动跳转至登录页面，提示“请先登录系统后再进行操作”；')
    p5_4 = doc.add_paragraph(style='No Indent')
    p5_4.add_run('• 会话过期：用户长时间未操作（超过30分钟），再次操作时提示“会话已过期，请重新登录”，自动跳转至登录页面；')
    p5_5 = doc.add_paragraph(style='No Indent')
    p5_5.add_run('• 数据为空：专病信息列表无数据时，展示“暂无专病信息，请先录入”，无查询结果时展示“未查询到相关数据”。')

    # ======================================
    # 六、附录（可选）
    # ======================================
    h1_6 = doc.add_paragraph('六、附录（可选）', style='Heading 1 Custom')
    # 6.1 术语说明
    h2_61 = doc.add_paragraph('6.1 术语说明', style='Heading 2 Custom')
    p61_1 = doc.add_paragraph(style='No Indent')
    p61_1.add_run('• 专病库：医疗领域中针对特定疾病建立的标准化信息数据库，包含疾病编码、名称、诊疗方案等核心信息；')
    p61_2 = doc.add_paragraph(style='No Indent')
    p61_2.add_run('• 科室管理员：各临床科室负责专病信息审核、管理的工作人员，拥有审核、编辑、删除的最高权限；')
    p61_3 = doc.add_paragraph(style='No Indent')
    p61_3.add_run('• 公开查询：未登录用户可访问的查询页面，仅展示审核通过的专病基础信息，无详细诊疗方案；')
    p61_4 = doc.add_paragraph(style='No Indent')
    p61_4.add_run('• 精准查询：根据唯一的专病编码或名称进行的精确匹配查询，仅返回一条结果。')
    # 6.2 补充说明
    h2_62 = doc.add_paragraph('6.2 补充说明', style='Heading 2 Custom')
    p62_1 = doc.add_paragraph(style='No Indent')
    p62_1.add_run('• 本版本需求中，“专病数据批量导入/导出功能”暂不开发，无需生成相关测试用例；')
    p62_2 = doc.add_paragraph(style='No Indent')
    p62_2.add_run('• 专病信息的“删除功能”仅科室管理员拥有权限，普通用户无该操作按钮，无需单独测试权限控制；')
    p62_3 = doc.add_paragraph(style='No Indent')
    p62_3.add_run('• 需求变更记录：2026-02-18 新增“专病信息录入后需管理员审核”规则，原无审核环节。')

    # ======================================
    # 七、格式编辑规范（必选）
    # ======================================
    h1_7 = doc.add_paragraph('七、格式编辑规范（必选，编辑时严格遵守）', style='Heading 1 Custom')
    p7_1 = doc.add_paragraph(style='No Indent')
    p7_1.add_run('• 标题层级：严格沿用本文档的一级（X.XXX）、二级（X.X XXX）、三级（X.X.X XXX）、四级（X.X.X.X XXX）标题，不可新增层级或修改层级名称；')
    p7_2 = doc.add_paragraph(style='No Indent')
    p7_2.add_run('• 文本格式：所有正文采用宋体、小四字体，行间距1.5倍，段落对齐方式：标题居中，正文左对齐；')
    p7_3 = doc.add_paragraph(style='No Indent')
    p7_3.add_run('• 表格规范：输入项/输出项等需结构化呈现的内容，必须用表格，表格无合并单元格、无边框装饰（默认无边框即可），字段名称不可修改；')
    p7_4 = doc.add_paragraph(style='No Indent')
    p7_4.add_run('• 冗余控制：每个模块、每个功能点的描述，仅保留核心信息，避免口语化、冗余化表述（如避免“这个功能主要是用来...”，直接用“功能用途：XXX”）；')
    p7_5 = doc.add_paragraph(style='No Indent')
    p7_5.add_run('• 一致性：同一术语、同一格式在全文中保持一致（如“必填”不可写为“必须填写”，“专病编码”不可写为“疾病编码”）；')
    p7_6 = doc.add_paragraph(style='No Indent')
    p7_6.add_run('• 避免插入：文档中不可插入图片、截图、图表、公式等，仅保留文本和表格，避免干扰通义千问提取需求。')

    # ======================================
    # 八、适配通义千问API的补充说明
    # ======================================
    h1_8 = doc.add_paragraph('八、适配通义千问API的补充说明', style='Heading 1 Custom')
    p8_1 = doc.add_paragraph(style='No Indent')
    p8_1.add_run('1. 按此格式编写的docx需求文档，可直接用于之前的Python代码（替换DOCX_PATH路径），无需修改Prompt，通义千问可精准提取“模块-功能-输入-输出-规则”，生成符合规范的测试用例；')
    p8_2 = doc.add_paragraph(style='No Indent')
    p8_2.add_run('2. 若需求变更，仅修改对应模块的“功能描述/输入项/输出项”等内容，保持格式不变，确保API调用时，大模型能快速识别变更后的需求；')
    p8_3 = doc.add_paragraph(style='No Indent')
    p8_3.add_run('3. 建议每个功能点独立成节，避免多个功能点合并编写，减少大模型提取信息时的混淆；')
    p8_4 = doc.add_paragraph(style='No Indent')
    p8_4.add_run('4. 专病库系统的新增功能需按本模版的结构化格式补充，确保通义千问API调用的一致性和准确性；')
    p8_5 = doc.add_paragraph(style='No Indent')
    p8_5.add_run('5. 表格中的输入项规则为通义千问生成用例的核心依据，请勿修改表格字段和格式。')

    # 保存文档
    doc.save('【专病库】需求规格说明书（测试用例生成专用）.docx')
    print('需求文档生成成功！文件名为：【专病库】需求规格说明书（测试用例生成专用）.docx')

if __name__ == '__main__':
    create_medical_disease_lib_docs()