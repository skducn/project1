# -*- coding: utf-8 -*-
# *****************************************************************
# Author     : John
# Date       : 2021-12-30
# Description: # python-docx 0.8.11
# 官网：https://www.osgeo.cn/python-docx/index.html
# python-docx添加和删除表格行、列 https://www.cnblogs.com/aademeng/articles/13056803.html
# python docx 使用总结 https://www.jianshu.com/p/7d2fcf976914
# pip install python-docx
# pip3 install pdf2docx
# pdf2docx https://blog.csdn.net/sinat_15136141/article/details/113620116
# *****************************************************************

"""
1 新建docx
2 替换docx中内容
3 docx转pdf
4 pdf转docx®®

"""
from docx import Document
from docx2pdf import convert
from pdf2docx import Converter

from PO.SysPO import *

Sys_PO = SysPO()


class DocxPO:
    def newDocx(self, filename):
        """
        1 新建docx
        :param filename:
        :return:
        """
        document = Document()
        document.add_heading("Document Title first", 0)

        p = document.add_paragraph("A plain paragraph having some ")
        p.add_run("bold").bold = True
        p.add_run(" and some first")
        p.add_run("italic.").italic = True

        document.add_heading("Heading, level 1 first", level=1)
        document.add_paragraph("Intense quote first", style="Intense Quote")

        document.add_paragraph(
            "first item in unordered list first", style="List Bullet"
        )
        document.add_paragraph("first item in ordered list first", style="List Number")

        # document.add_picture('test.jpg', width=Inches(1.25))

        records = (
            (3, "101", "Spam"),
            (7, "422 first", "Eggs"),
            (4, "631", "Spam, spam, eggs, and spam first"),
        )

        table = document.add_table(rows=1, cols=3)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Qty"
        hdr_cells[1].text = "Id first"
        hdr_cells[2].text = "Desc"
        for qty, id, desc in records:
            row_cells = table.add_row().cells
            row_cells[0].text = str(qty)
            row_cells[1].text = id
            row_cells[2].text = desc

        document.add_page_break()

        document.save(filename)

    def replaceText(self, filename, org_content, tag_content):
        """
        2 替换docx中内容
        :param filename:
        :return:
        """

        document = Document(filename)

        for para in document.paragraphs:
            # print(para.text)
            if org_content in para.text:
                for run in para.runs:
                    if org_content in run.text:
                        run.text = run.text.replace(org_content, tag_content)

        for t in document.tables:
            for i in range(len(t.rows)):
                for j in range(len(t.columns)):
                    # print(t.cell(i, j).text)
                    if org_content in t.cell(i, j).text:
                        t.cell(i, j).text = t.cell(i, j).text.replace(
                            org_content, tag_content
                        )

        document.save(filename)

    def docx2pdf(self, filename):
        """
        3 docx转pdf
        :param filename:
        :return:
        """
        convert(filename)

    def pdf2docx(self, pdf_file, docx_file):
        """
        4 pdf转docx
        :return:
        """

        cv = Converter(pdf_file)
        cv.convert(docx_file, start=0, end=None)
        cv.close()


if __name__ == "__main__":

    Sys_PO.killPid("WINWORD.EXE")

    Docx_PO = DocxPO()

    # print("1 新建docx".center(100, "-"))
    # Docx_PO.newDocx('./data/test1.docx')
    #
    # print("2 替换docx中内容".center(100, "-"))
    # Docx_PO.replaceText('./data/test1.docx', 'first', '百度')  # 将文档中 first字样替换成百度
    #
    # print("3 docx转换成pdf".center(100, "-"))
    # Docx_PO.docx2pdf('./data/test1.docx')
    #

    print("4 pdf转换成docx".center(100, "-"))
    Docx_PO.pdf2docx(
        "/Users/linghuchong/Downloads/51Testing_wenzhang65_2.pdf",
        "/Users/linghuchong/Downloads/51Testing_wenzhang65_2.docx",
    )
